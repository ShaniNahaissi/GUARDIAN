import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_paragraph(doc, text, style='Normal', space_after=6, line_spacing=1.15, is_centered=False):
    p = doc.add_paragraph(style=style)
    if is_centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    # Robust inline formatting: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x7F, 0x1D, 0x1D)
        else:
            p.add_run(part)
    return p

def add_image_to_docx(doc, image_name):
    paths_to_try = [
        os.path.join("metrics", image_name),
        os.path.join("frontend", "frontend_definition", image_name),
        os.path.join("temporal_training", "figures", image_name),
        image_name
    ]
    found_path = None
    for p in paths_to_try:
        if os.path.exists(p):
            found_path = p
            break
    
    if found_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(found_path, width=Inches(5.8))
        print(f"Embedded image: {found_path}")
        return True
    else:
        print(f"WARNING: Image not found: {image_name}")
        return False

def extract_missing_figures():
    from pathlib import Path
    figures_dir = Path("temporal_training/figures")
    notebook_path = Path("temporal_training/temporal_training.ipynb")
    
    needed_files = ["learning_curves.png", "confusion_matrix.png", "roc_pr_curves.png", "class_balance.png"]
    missing = [f for f in needed_files if not (figures_dir / f).exists()]
    
    if not missing:
        return
        
    if not notebook_path.exists():
        print(f"WARNING: Needed figures {missing} are missing and {notebook_path} does not exist to extract them.")
        return
        
    print(f"Auto-extracting missing figures {missing} from {notebook_path}...")
    import json
    import base64
    try:
        with open(notebook_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            
        figures_dir.mkdir(parents=True, exist_ok=True)
        extracted = 0
        for cell in data.get("cells", []):
            if "outputs" in cell:
                for out in cell["outputs"]:
                    if "data" in out and "image/png" in out["data"]:
                        img_data = out["data"]["image/png"]
                        img_bytes = base64.b64decode(img_data.strip())
                        
                        source = "".join(cell.get("source", []))
                        filename = None
                        if "learning_curves.png" in source or "train_loss" in source:
                            filename = "learning_curves.png"
                        elif "confusion_matrix.png" in source or "ConfusionMatrixDisplay" in source:
                            filename = "confusion_matrix.png"
                        elif "roc_pr_curves.png" in source or "roc_curve" in source:
                            filename = "roc_pr_curves.png"
                        elif "class_balance.png" in source or "sequences_saved" in source:
                            filename = "class_balance.png"
                            
                        if filename and filename in missing:
                            dest_path = figures_dir / filename
                            dest_path.write_bytes(img_bytes)
                            print(f"Auto-extracted: {filename}")
                            extracted += 1
        print(f"Auto-extracted {extracted} plots successfully.")
    except Exception as e:
        print(f"WARNING: Failed to auto-extract figures from notebook: {e}")

def convert_md_to_docx(md_path, output_docx_path):
    extract_missing_figures()
    doc = Document()

    # Set page margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure Default Font
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_code_block():
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        
        full_code = "".join(code_lines).rstrip()
        run = p.add_run(full_code)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        code_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        # Parse markdown table
        rows_data = []
        for l in table_lines:
            if '---' in l:
                continue
            cols = [c.strip() for c in l.strip().split('|')[1:-1]]
            if cols:
                rows_data.append(cols)
        
        if rows_data:
            table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            for r_idx, row_cols in enumerate(rows_data):
                for c_idx, text in enumerate(row_cols):
                    if c_idx < len(table.columns):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
                        
                        pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
                        parts = pattern.split(text)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            elif part.startswith('*') and part.endswith('*'):
                                run = p.add_run(part[1:-1])
                                run.italic = True
                            elif part.startswith('`') and part.endswith('`'):
                                run = p.add_run(part[1:-1])
                                run.font.name = 'Consolas'
                                run.font.size = Pt(9.0)
                            else:
                                p.add_run(part)
                                
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
                            set_cell_background(cell, "E5E7EB")
                        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            doc.add_paragraph() # spacing after table
        table_lines = []

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    is_centered = False

    for l in lines:
        line = l.rstrip('\r\n')

        if "End of Final Project Book" in line:
            continue

        if "<!-- pagebreak -->" in line or "\\pagebreak" in line:
            doc.add_page_break()
            continue

        if '<div align="center">' in line or "<div align='center'>" in line:
            is_centered = True
            continue

        if '</div>' in line:
            is_centered = False
            continue
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line + '\n')
            continue

        # Table lines
        if line.strip().startswith('|') and line.strip().endswith('|'):
            in_table = True
            table_lines.append(line)
            continue
        elif in_table:
            flush_table()
            in_table = False

        # Horizontal rule
        if line.strip() == '---' or line.strip() == '***':
            doc.add_paragraph()
            continue

        # Headings
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=0)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            continue
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=1)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            continue
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=2)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            continue
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=3)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            continue
        elif line.startswith('##### '):
            p = doc.add_heading(line[6:].strip(), level=4)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            continue

        # Bullet lists
        if line.strip().startswith('* ') or line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            add_styled_paragraph(doc, text, style='List Bullet', space_after=3)
            continue

        # Ordered lists
        match = re.match(r'^\d+\.\s+(.*)', line.strip())
        if match:
            text = match.group(1)
            add_styled_paragraph(doc, text, style='List Number', space_after=3)
            continue

        # Normal text
        if line.strip():
            if "Note: Manually insert" in line:
                img_matches = re.findall(r'\[(.*?\.png)\]', line)
                allowed_embeds = [img for img in img_matches if img in {"main_page.png", "camera_view.png", "current_metrics1.png", "current_metrics2.png", "learning_curves.png", "confusion_matrix.png"}]
                if allowed_embeds:
                    for img_name in allowed_embeds:
                        add_image_to_docx(doc, img_name)
                    # Extract the description part after "here. "
                    desc_parts = line.split("here. ")
                    if len(desc_parts) > 1:
                        desc_text = desc_parts[1].rstrip('* ')
                        p_desc = doc.add_paragraph()
                        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_desc.paragraph_format.space_before = Pt(4)
                        p_desc.paragraph_format.space_after = Pt(12)
                        run = p_desc.add_run(desc_text)
                        run.italic = True
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(0x52, 0x52, 0x52)
                    continue
            add_styled_paragraph(doc, line.strip(), style='Normal', space_after=6, is_centered=is_centered)

    try:
        doc.save(output_docx_path)
        print(f"Successfully generated Word document: {output_docx_path}")
    except PermissionError:
        fallback_path = output_docx_path.replace(".docx", "_New.docx")
        print(f"WARNING: Permission denied writing to {output_docx_path}. The file might be open in another application.")
        print(f"Attempting to write to fallback file: {fallback_path}")
        doc.save(fallback_path)
        print(f"Successfully generated fallback Word document: {fallback_path}")

if __name__ == '__main__':
    md_path = os.path.join(os.path.dirname(__file__), 'FINAL_PROJECT_BOOK.md')
    out_path = os.path.join(os.path.dirname(__file__), 'GUARDIAN_Final_Project_Book_New.docx')
    convert_md_to_docx(md_path, out_path)
