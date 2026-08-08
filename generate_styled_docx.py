import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- Color Palette Constants (Academic Navy & Slate) ---
COLOR_PRIMARY = RGBColor(0x1B, 0x36, 0x5D)    # Deep Navy for Title & Main Headings
COLOR_SECONDARY = RGBColor(0x2B, 0x4C, 0x7E)  # Slate Navy for Subheadings
COLOR_TERTIARY = RGBColor(0x3A, 0x5A, 0x80)   # Steel Blue for Minor Headings
COLOR_BODY = RGBColor(0x26, 0x26, 0x26)       # Dark Charcoal for Body Text
COLOR_MUTED = RGBColor(0x5A, 0x6B, 0x7C)      # Cool Gray for Captions & Headers
HEX_PRIMARY = "1B365D"
HEX_ZEBRA = "F8FAFC"
HEX_BORDER = "E2E8F0"
HEX_CODE_BG = "F4F6F9"
HEX_CODE_BORDER = "1B365D"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color_hex="1B365D", sz="24"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

def set_table_academic_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="12" w:space="0" w:color="{HEX_PRIMARY}"/>
            <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{HEX_PRIMARY}"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_header_footer(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("GUARDIAN: Near Real-Time Threat & Behavioral Detection System")
    hrun.font.name = 'Calibri'
    hrun.font.size = Pt(8.5)
    hrun.font.italic = True
    hrun.font.color.rgb = COLOR_MUTED

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    frun = fp.add_run("Page ")
    frun.font.name = 'Calibri'
    frun.font.size = Pt(9)
    frun.font.color.rgb = COLOR_MUTED

    fldSimple1 = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
    fp._p.append(fldSimple1)

    frun2 = fp.add_run(" of ")
    frun2.font.name = 'Calibri'
    frun2.font.size = Pt(9)
    frun2.font.color.rgb = COLOR_MUTED

    fldSimple2 = parse_xml(r'<w:fldSimple %s w:instr="NUMPAGES"/>' % nsdecls('w'))
    fp._p.append(fldSimple2)

def add_styled_paragraph(doc, text, style='Normal', space_before=0, space_after=5, line_spacing=1.18, is_centered=False):
    p = doc.add_paragraph(style=style)
    if is_centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_BODY
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_BODY
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B) # Crimson code pill
        else:
            run = p.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_BODY
    return p

def add_image_to_docx(doc, image_name):
    paths_to_try = [
        os.path.join("metrics", image_name),
        os.path.join("frontend", "frontend_definition", image_name),
        os.path.join("temporal_training", "figures", image_name),
        image_name
    ]
    found_path = None
    for p in paths_to_try:
        if os.path.exists(p):
            found_path = p
            break
    
    if found_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(found_path, width=Inches(5.8))
        print(f"Embedded image: {found_path}")
        return True
    else:
        print(f"WARNING: Image not found: {image_name}")
        return False

def extract_missing_figures():
    figures_dir = Path("temporal_training/figures")
    notebook_path = Path("temporal_training/temporal_training.ipynb")
    
    needed_files = ["learning_curves.png", "confusion_matrix.png", "roc_pr_curves.png", "class_balance.png"]
    missing = [f for f in needed_files if not (figures_dir / f).exists()]
    
    if not missing:
        return
        
    if not notebook_path.exists():
        print(f"WARNING: Needed figures {missing} are missing and {notebook_path} does not exist to extract them.")
        return
        
    print(f"Auto-extracting missing figures {missing} from {notebook_path}...")
    import json
    import base64
    try:
        with open(notebook_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            
        figures_dir.mkdir(parents=True, exist_ok=True)
        extracted = 0
        for cell in data.get("cells", []):
            if "outputs" in cell:
                for out in cell["outputs"]:
                    if "data" in out and "image/png" in out["data"]:
                        img_data = out["data"]["image/png"]
                        img_bytes = base64.b64decode(img_data.strip())
                        
                        source = "".join(cell.get("source", []))
                        filename = None
                        if "learning_curves.png" in source or "train_loss" in source:
                            filename = "learning_curves.png"
                        elif "confusion_matrix.png" in source or "ConfusionMatrixDisplay" in source:
                            filename = "confusion_matrix.png"
                        elif "roc_pr_curves.png" in source or "roc_curve" in source:
                            filename = "roc_pr_curves.png"
                        elif "class_balance.png" in source or "sequences_saved" in source:
                            filename = "class_balance.png"
                            
                        if filename and filename in missing:
                            dest_path = figures_dir / filename
                            dest_path.write_bytes(img_bytes)
                            print(f"Auto-extracted: {filename}")
                            extracted += 1
        print(f"Auto-extracted {extracted} plots successfully.")
    except Exception as e:
        print(f"WARNING: Failed to auto-extract figures from notebook: {e}")

def convert_md_to_styled_docx(md_path, output_docx_path):
    extract_missing_figures()
    doc = Document()

    # Set page margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure Header & Footer
    add_header_footer(doc)

    # Configure Default Base Style
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = COLOR_BODY

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_code_block():
        nonlocal code_lines
        if not code_lines:
            return
        
        # Render Code Block as a single-cell callout container table
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        
        set_cell_background(cell, HEX_CODE_BG)
        set_cell_left_border(cell, color_hex=HEX_CODE_BORDER, sz="24") # 3pt left accent bar
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        
        full_code = "".join(code_lines).rstrip()
        run = p.add_run(full_code)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        # Spacing after code block
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(0)
        p_space.paragraph_format.space_after = Pt(6)
        
        code_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        
        rows_data = []
        for l in table_lines:
            if '---' in l:
                continue
            cols = [c.strip() for c in l.strip().split('|')[1:-1]]
            if cols:
                rows_data.append(cols)
        
        if rows_data:
            table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_academic_borders(table)
            
            for r_idx, row_cols in enumerate(rows_data):
                is_header = (r_idx == 0)
                is_zebra = (r_idx % 2 == 1 and not is_header)
                
                for c_idx, text in enumerate(row_cols):
                    if c_idx < len(table.columns):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.line_spacing = 1.05
                        
                        pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
                        parts = pattern.split(text)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            elif part.startswith('*') and part.endswith('*'):
                                run = p.add_run(part[1:-1])
                                run.italic = True
                            elif part.startswith('`') and part.endswith('`'):
                                run = p.add_run(part[1:-1])
                                run.font.name = 'Consolas'
                                run.font.size = Pt(8.5)
                            else:
                                run = p.add_run(part)
                            
                            run.font.name = 'Calibri'
                            if is_header:
                                run.bold = True
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            else:
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = COLOR_BODY

                        if is_header:
                            set_cell_background(cell, HEX_PRIMARY)
                        elif is_zebra:
                            set_cell_background(cell, HEX_ZEBRA)
                            
                        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(0)
            p_space.paragraph_format.space_after = Pt(6)
            
        table_lines = []

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    is_centered = False

    for l in lines:
        line = l.rstrip('\r\n')

        if "End of Final Project Book" in line:
            continue

        if "<!-- pagebreak -->" in line or "\\pagebreak" in line:
            doc.add_page_break()
            continue

        if '<div align="center">' in line or "<div align='center'>" in line:
            is_centered = True
            continue

        if '</div>' in line:
            is_centered = False
            continue
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line + '\n')
            continue

        # Table lines
        if line.strip().startswith('|') and line.strip().endswith('|'):
            in_table = True
            table_lines.append(line)
            continue
        elif in_table:
            flush_table()
            in_table = False

        # Horizontal rule
        if line.strip() == '---' or line.strip() == '***':
            p_hr = doc.add_paragraph()
            p_hr.paragraph_format.space_before = Pt(8)
            p_hr.paragraph_format.space_after = Pt(8)
            continue

        # Headings (Strictly preserving original text and levels)
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=0)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(8)
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(22)
                r.font.bold = True
                r.font.color.rgb = COLOR_PRIMARY
            continue
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=1)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(16)
                r.font.bold = True
                r.font.color.rgb = COLOR_PRIMARY
            continue
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=2)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(13)
                r.font.bold = True
                r.font.color.rgb = COLOR_SECONDARY
            continue
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=3)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(11.5)
                r.font.bold = True
                r.font.color.rgb = COLOR_TERTIARY
            continue
        elif line.startswith('##### '):
            p = doc.add_heading(line[6:].strip(), level=4)
            if is_centered:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = COLOR_TERTIARY
            continue

        # Bullet lists
        if line.strip().startswith('* ') or line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            add_styled_paragraph(doc, text, style='List Bullet', space_after=3, line_spacing=1.15)
            continue

        # Ordered lists
        match = re.match(r'^\d+\.\s+(.*)', line.strip())
        if match:
            text = match.group(1)
            add_styled_paragraph(doc, text, style='List Number', space_after=3, line_spacing=1.15)
            continue

        # Normal text & Image Insert Notes
        if line.strip():
            if "Note: Manually insert" in line:
                img_matches = re.findall(r'\[(.*?\.png)\]', line)
                allowed_embeds = [img for img in img_matches if img in {"main_page.png", "camera_view.png", "current_metrics1.png", "current_metrics2.png", "learning_curves.png", "confusion_matrix.png"}]
                if allowed_embeds:
                    for img_name in allowed_embeds:
                        add_image_to_docx(doc, img_name)
                    desc_parts = line.split("here. ")
                    if len(desc_parts) > 1:
                        desc_text = desc_parts[1].rstrip('* ')
                        p_desc = doc.add_paragraph()
                        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_desc.paragraph_format.space_before = Pt(4)
                        p_desc.paragraph_format.space_after = Pt(12)
                        run = p_desc.add_run(desc_text)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9.5)
                        run.font.italic = True
                        run.font.color.rgb = COLOR_MUTED
                    continue
            add_styled_paragraph(doc, line.strip(), style='Normal', space_after=5, line_spacing=1.18, is_centered=is_centered)

    try:
        doc.save(output_docx_path)
        print(f"Successfully generated Styled Word document: {output_docx_path}")
    except PermissionError:
        fallback_path = output_docx_path.replace(".docx", "_New.docx")
        print(f"WARNING: Permission denied writing to {output_docx_path}. Writing to fallback: {fallback_path}")
        doc.save(fallback_path)
        print(f"Successfully generated fallback Styled Word document: {fallback_path}")

if __name__ == '__main__':
    md_path = os.path.join(os.path.dirname(__file__), 'FINAL_PROJECT_BOOK.md')
    out_path = os.path.join(os.path.dirname(__file__), 'GUARDIAN_Final_Project_Book_Styled.docx')
    convert_md_to_styled_docx(md_path, out_path)
